#!/usr/bin/env python3
"""
生成 Word 格式的学生成绩报告并发送邮件
"""

import sys
import json
import smtplib
from pathlib import Path
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.utils import formataddr
from datetime import datetime

# 添加本地库路径
lib_path = Path(__file__).parent.parent.parent / "lib"
sys.path.insert(0, str(lib_path))

# 导入 python-docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 导入飞书工具模块
sys.path.insert(0, str(Path(__file__).parent))
from feishu_utils import get_access_token, read_sheet_range, parse_table_data

def get_student_data(student_name, spreadsheet_token, sheets_config):
    """获取学生成绩数据"""
    token = get_access_token()
    if not token:
        return None
    
    all_exams = []
    for exam_date, sheet_id in sheets_config.items():
        values = read_sheet_range(token, spreadsheet_token, sheet_id, "A1:P50")
        records = parse_table_data(values, key_field="学号")
        
        for r in records:
            if r.get("学生姓名") == student_name or r.get("姓名") == student_name:
                r["考试日期"] = exam_date
                all_exams.append(r)
                break
    
    return all_exams

def generate_word_report(student_name, exams, output_path):
    """生成 Word 格式成绩报告"""
    doc = Document()
    
    # 标题
    title = doc.add_heading(f'{student_name}同学考试成绩报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    latest = exams[0]
    doc.add_heading('一、基本信息', level=1)
    p = doc.add_paragraph()
    p.add_run(f'姓名：{student_name}\n')
    p.add_run(f'班级：{latest.get("班级", "N/A")}\n')
    p.add_run(f'学号：{latest.get("学号", "N/A")}\n')
    p.add_run(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}\n')
    
    # 最近考试概览
    doc.add_heading('二、最近考试概览', level=1)
    
    table = doc.add_table(rows=len(exams)+1, cols=5)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['考试日期', '年级排名', '总分', '平均分', '变化']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # 数据行
    for i, exam in enumerate(exams):
        row = table.rows[i+1]
        row.cells[0].text = exam.get("考试日期", "N/A")
        row.cells[1].text = str(exam.get("年级排名", "N/A"))
        row.cells[2].text = str(exam.get("总分", "N/A"))
        row.cells[3].text = str(exam.get("平均分", "N/A"))
        
        # 计算变化
        if i < len(exams) - 1:
            prev = exams[i+1]
            prev_rank = int(prev.get("年级排名", 0) or 0)
            curr_rank = int(exam.get("年级排名", 0) or 0)
            prev_total = int(prev.get("总分", 0) or 0)
            curr_total = int(exam.get("总分", 0) or 0)
            
            rank_change = prev_rank - curr_rank
            total_change = curr_total - prev_total
            
            if rank_change > 0:
                row.cells[4].text = f'排名↑{rank_change}名, 总分+{total_change}分'
            elif rank_change < 0:
                row.cells[4].text = f'排名↓{abs(rank_change)}名, 总分{total_change}分'
            else:
                row.cells[4].text = f'排名持平, 总分+{total_change}分'
        else:
            row.cells[4].text = '-'
    
    doc.add_paragraph()  # 空行
    
    # 各科成绩对比
    doc.add_heading('三、各科成绩对比', level=1)
    
    subjects = ["语文", "数学", "英语", "政治", "历史", "地理", "生物"]
    
    table2 = doc.add_table(rows=len(subjects)+1, cols=len(exams)+2)
    table2.style = 'Table Grid'
    
    # 表头
    table2.rows[0].cells[0].text = '科目'
    for i, exam in enumerate(exams):
        table2.rows[0].cells[i+1].text = exam.get("考试日期", "N/A")
    table2.rows[0].cells[-1].text = '变化'
    
    # 数据行
    for j, subj in enumerate(subjects):
        row = table2.rows[j+1]
        row.cells[0].text = subj
        
        scores = []
        for i, exam in enumerate(exams):
            score = exam.get(subj, "N/A")
            scores.append(int(score) if score and score != "N/A" else 0)
            row.cells[i+1].text = str(score)
        
        # 计算变化
        if len(scores) >= 2 and scores[0] and scores[1]:
            change = scores[0] - scores[1]
            row.cells[-1].text = f'+{change}' if change > 0 else str(change)
        else:
            row.cells[-1].text = '-'
    
    doc.add_paragraph()
    
    # 学科分析
    doc.add_heading('四、学科分析', level=1)
    
    latest_scores = {}
    for subj in subjects:
        score = latest.get(subj)
        if score and score != "N/A":
            latest_scores[subj] = int(score)
    
    if latest_scores:
        sorted_scores = sorted(latest_scores.items(), key=lambda x: x[1], reverse=True)
        p = doc.add_paragraph()
        p.add_run('优势学科：').bold = True
        p.add_run(', '.join([f'{s[0]}({s[1]}分)' for s in sorted_scores[:3]]) + '\n')
        p.add_run('薄弱学科：').bold = True
        p.add_run(', '.join([f'{s[0]}({s[1]}分)' for s in sorted_scores[-3:]]) + '\n')
    
    # 教师评语
    doc.add_heading('五、教师评语', level=1)
    
    latest_rank = int(latest.get("年级排名", 0) or 0)
    p = doc.add_paragraph()
    
    if latest_rank == 1:
        p.add_run(f'🌟 {student_name}同学本次考试年级第一，表现优异！\n')
        p.add_run('各科成绩均衡，继续保持，争取更大突破！')
    elif latest_rank <= 5:
        p.add_run(f'👍 {student_name}同学本次考试年级前五，表现优秀！\n')
        if len(exams) > 1:
            prev_rank = int(exams[1].get("年级排名", 0) or 0)
            if latest_rank < prev_rank:
                p.add_run('排名有所提升，进步明显，继续保持！')
    elif latest_rank <= 10:
        p.add_run(f'📈 {student_name}同学本次考试年级前十，表现良好。\n')
        p.add_run('继续努力，争取进入前五！')
    else:
        p.add_run(f'💪 {student_name}同学要继续努力，争取更大进步！')
    
    # 家校沟通
    doc.add_heading('六、家校沟通', level=1)
    p = doc.add_paragraph()
    p.add_run('如需详细了解孩子学习情况，欢迎与老师联系沟通。\n')
    p.add_run('让我们共同关注孩子的成长！')
    
    # 保存
    doc.save(output_path)
    return output_path

def send_email_with_attachment(to_email, subject, body, attachment_path, smtp_config):
    """发送带附件的邮件"""
    msg = MIMEMultipart()
    msg['From'] = formataddr((smtp_config['sender_name'], smtp_config['email']))
    msg['To'] = to_email
    msg['Subject'] = subject
    
    # 邮件正文
    msg.attach(MIMEText(body, 'plain', 'utf-8'))
    
    # 附件
    with open(attachment_path, 'rb') as f:
        attachment = MIMEApplication(f.read())
        attachment.add_header('Content-Disposition', 'attachment', 
                            filename=Path(attachment_path).name)
        msg.attach(attachment)
    
    # 发送
    with smtplib.SMTP_SSL(smtp_config['server'], smtp_config['port']) as server:
        server.login(smtp_config['email'], smtp_config['auth_code'])
        server.send_message(msg)
    
    return True

def main():
    if len(sys.argv) < 2:
        print("用法: python send_word_report.py <学生姓名> [家长邮箱]")
        sys.exit(1)
    
    student_name = sys.argv[1]
    parent_email = sys.argv[2] if len(sys.argv) > 2 else None
    
    # 读取配置
    config_path = Path(__file__).parent.parent.parent.parent / "config.json"
    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)
    
    smtp_config = config.get("smtp", {})
    
    if not smtp_config.get("email"):
        print("❌ SMTP 未配置")
        sys.exit(1)
    
    # 成绩汇总表配置
    spreadsheet_token = "SXlnsuW9rhhbQPtjHQKciZZXnUf"
    sheets_config = {
        "2025-11-12": "CNHOVc",
        "2025-10-25": "iintOO",
        "2025-10-12": "9dd137"
    }
    
    # 获取学生数据
    print(f"正在获取 {student_name} 的成绩数据...")
    exams = get_student_data(student_name, spreadsheet_token, sheets_config)
    
    if not exams:
        print(f"❌ 未找到 {student_name} 的成绩数据")
        sys.exit(1)
    
    print(f"✅ 找到 {len(exams)} 次考试成绩")
    
    # 生成 Word 报告
    output_dir = Path(__file__).parent.parent.parent.parent / "data" / "reports"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{student_name}_成绩报告.docx"
    
    print(f"\n正在生成 Word 报告...")
    generate_word_report(student_name, exams, output_path)
    print(f"✅ Word 报告已生成: {output_path}")
    
    # 发送邮件
    if parent_email:
        print(f"\n正在发送邮件给 {parent_email}...")
        
        latest = exams[0]
        subject = f"【成绩报告】{student_name}同学最近考试成绩报告"
        body = f"""
尊敬的家长，您好！

附件是 {student_name} 同学最近的考试成绩报告（Word格式）。

【最近考试概况】
考试日期：{latest.get('考试日期', 'N/A')}
年级排名：第{latest.get('年级排名', 'N/A')}名
总分：{latest.get('总分', 'N/A')}分
平均分：{latest.get('平均分', 'N/A')}分

详细内容请查看附件。

如有疑问，欢迎与老师联系沟通。

祝好！
{smtp_config.get('sender_name', '老师')}
{datetime.now().strftime('%Y年%m月%d日')}
"""
        
        send_email_with_attachment(parent_email, subject, body, output_path, smtp_config)
        print(f"✅ 邮件已发送成功（附带 Word 报告）！")
    else:
        print(f"\n⚠️ 未提供家长邮箱")
        print(f"Word 报告已保存到: {output_path}")

if __name__ == "__main__":
    main()
